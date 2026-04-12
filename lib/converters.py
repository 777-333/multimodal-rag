"""Text extractors for Office file formats."""
from __future__ import annotations
import io
import csv as csv_module


def extract_docx_text(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)


def extract_xlsx_text(file_bytes: bytes) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    parts = []
    for sheet in wb.worksheets:
        parts.append(f"=== {sheet.title} ===")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            line = "\t".join(cells)
            if line.strip():
                parts.append(line)
    return "\n".join(parts)


def extract_csv_text(file_bytes: bytes) -> str:
    text = file_bytes.decode("utf-8", errors="replace")
    reader = csv_module.reader(io.StringIO(text))
    return "\n".join("\t".join(row) for row in reader if any(row))
